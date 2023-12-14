from docx import Document
from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert #워드를 pdf로 변환하기 위한 라이브러리를 불러옵니다.

load_wd=load_workbook(r"12. 엑셀의 정보를 불러와 수료증 자동생성\수료증명단.xlsx")
load_ws=load_wd.active #8~21: 엑셀에서 값을 읽습니다.

name_list=[]
birthday_list=[]
ho_list=[]
for i in range(1,load_ws.max_row+1):
  name_list.append(load_ws.cell(i,1).value)
  birthday_list.append(load_ws.cell(i,2).value)
  ho_list.append(load_ws.cell(i,3).value)

print(name_list)
print(birthday_list)
print(ho_list)

for i in range(len(name_list)):#21~95:엑셀에서 읽은 이름 리스트의 길이만큼 반복합니다. 즉, 수료증 수만큼 반복합니다.
  doc=docx.Document(r'12. 엑셀의 정보를 불러와 수료증 자동생성\수료증양식.docx')
  style=doc.styles['Normal']
  style.font.name='나눔고딕'
  style._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
  run.font.size=docx.shared.Pt(20)

  para=doc.add_paragraph()
  run=para.add_run('\n\n')
  run=para.add_run(' 제'+ho_list[i]+' 호\n')#:ho_list로 엑셀에서 읽은 호(번호)로 입력합니다.
  run.font.name='나눔고딕'
  run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
  run.font.size=docx.shared.Pt(20)

  para=doc.add_paragraph()
  run=para.add_run('\n\n')
  run=para.add_run('수 료 증')
  run.font.name='나눔고딕'
  run.bold=True
  run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
  run.font.size=docx.shared.Pt(40)
  para.alignment=WD_ALIGN_PARAGRAPH.CENTER

  para=doc.add_paragraph()
  run=para.add_run('\n\n')
  run=para.add_run(' 성명:'+name_list[i]+'\n')#:이름부분은 엑셀에서 읽은 name_list로 합니다.
  run.font.name='나눔고딕'
  run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
  run.font.size=docx.shared.Pt(20)
  run-para.add_run(' 생 년 월 일: '+birthday_list[i]+'\n')#:생일부분은 엑셀에서 읽은 birthday_list로 합니다.
  run.font.name='나눔고딕'
  run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
  run.font.size=docx.shared.Pt(20)
  run=para.add_run(' 교 육 과 정: 파이썬과 40개의 작품들\n')
  run.font.name='나눔고딕'
  run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
  run.font.size=docx.shared.Pt(20)
  run=para.add_run(' 교 육 날 짜: 2021.08.05~2021.09.09\n')
  run.font.name='나눔고딕'
  run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
  run.font.size=docx.shared.Pt(20)

  para=doc.add_paragraph()
  run=para.add_run('\n\n')
  run=para.add_run(' 위 사람은 파이썬과 40개의 작품들 교육과정을\n')
  run.font.name='나눔고딕'
  run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
  run.font.size=docx.shared.Pt(20)
  run=para.add_run(' 이수하였으므로 이 증서를 수여합니다.\n')
  run.font.name='나눔고딕'
  run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
  run.font.size=docx.shared.Pt(20)

  para=doc.add_paragraph()
  run=para.add_run('\n\n\n')
  run=para.add_run('2021.09.19')
  run.font.name='나눔고딕'
  run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
  run.font.size=docx.shared.Pt(20)
  para.alignment=WD_ALIGN_PARAGRAPH.CENTER

  para=doc.add_paragraph()
  run=para.add_run('\n\n\n')
  run=para.add_run('파이썬교육기관장')
  run.font.name='나눔고딕'
  run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
  run.font.size=docx.shared.Pt(20)
  para.alignment=WD_ALIGN_PARAGRAPH.CENTER

  doc.save('12. 엑셀의 정보를 불러와 수료증 자동생성\\'+name_list[i]+'.docx')
  convert('12. 엑셀의 정보를 불러와 수료증 자동생성\\'+name_list[i]+'.docx', #:이름으로 워드파일을 생성합니다.
          '12. 엑셀의 정보를 불러와 수료증 자동생성\\'+name_list[i]+'.pdf')#워드파일을 읽어 pdf로 변환합니다. 한 줄의 코드나 코드가 길어 두 줄로 표현하였습니다.